"""
IntelliLaw — FastAPI Backend
Offline AI-Native Legal Operating System for Africa
"""
import os
import sys
import json
import logging
import shutil
import uuid
from datetime import datetime, date, timedelta
from pathlib import Path
from typing import List, Optional, Dict, Any
from contextlib import asynccontextmanager

from fastapi import (
    FastAPI, Depends, HTTPException, UploadFile, File,
    Form, BackgroundTasks, status
)
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from sqlalchemy import func, desc, or_
from pydantic import BaseModel, EmailStr

from app.database import get_db, init_db, get_engine, get_database_info
from app.models import User, AISession, Activity, SystemSettings
from app.legal_models import (
    Client, Matter, MatterNote, Hearing, Task, LegalDocument,
    TimeEntry, Invoice, InvoiceLineItem, Payment,
    TrustAccount, TrustTransaction, Disbursement,
    ResearchSession, LegalPrecedent, Conflict,
    ConflictStatus, ConflictRiskLevel
)
from app.auth import (
    get_current_active_user, require_admin, get_password_hash,
    authenticate_user, create_access_token, ACCESS_TOKEN_EXPIRE_MINUTES
)
from app.demo_mode import smart_generate, get_demo_status
from app.legal_prompts import (
    get_legal_prompt, build_draft_contract_prompt,
    build_draft_letter_prompt, build_summarize_prompt, build_research_prompt
)
from app.conflict_service import search_potential_conflicts

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────
# LIFESPAN
# ─────────────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("IntelliLaw starting up...")

    # 1. Database
    try:
        init_db()
        logger.info("Database initialised")
    except Exception as e:
        logger.error(f"DB init error: {e}")

    # 2. Default admin (created once on first run)
    try:
        from app.database import get_session_local
        from app.auth import get_password_hash
        SessionLocal = get_session_local()
        db = SessionLocal()
        if db.query(User).count() == 0:
            admin_username = os.getenv("DEFAULT_ADMIN_USERNAME", "admin")
            admin_email    = os.getenv("DEFAULT_ADMIN_EMAIL", "admin@intellilaw.local")
            admin_password = os.getenv("DEFAULT_ADMIN_PASSWORD", "Admin123!")
            admin = User(
                username=admin_username,
                email=admin_email,
                hashed_password=get_password_hash(admin_password),
                full_name="System Administrator",
                title="Administrator",
                initials="SA",
                is_admin=True,
                is_active=True,
                user_role="admin",
            )
            db.add(admin)
            db.commit()
            logger.info(f"Default admin created — username: {admin_username}")
        else:
            logger.info(f"Users already exist — skipping default admin creation")
        db.close()
    except Exception as e:
        logger.warning(f"Admin creation warning: {e}")

    # 3. AI model (skip in demo mode)
    try:
        from app.demo_mode import is_demo_mode_available
        if not is_demo_mode_available():
            from app.local_llm_service import load_model
            import threading
            threading.Thread(target=load_model, daemon=True).start()
            logger.info("Local AI model loading in background…")
        else:
            logger.info("Demo mode active — skipping local model preload")
    except Exception as e:
        logger.warning(f"AI pre-load skipped: {e}")

    yield
    logger.info("IntelliLaw shutting down.")


# ─────────────────────────────────────────────────────────────
# APP
# ─────────────────────────────────────────────────────────────

app = FastAPI(
    title="IntelliLaw API",
    description="Offline AI-Native Legal Operating System",
    version="1.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ORIGINS", "*").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─────────────────────────────────────────────────────────────
# PYDANTIC SCHEMAS
# ─────────────────────────────────────────────────────────────

class LoginRequest(BaseModel):
    username: str
    password: str

class CreateUserRequest(BaseModel):
    username: str
    email: str
    password: str
    full_name: Optional[str] = None
    title: Optional[str] = None
    user_role: str = "attorney"

class ClientCreate(BaseModel):
    client_type: str = "individual"
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    company_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    id_number: Optional[str] = None
    address_line1: Optional[str] = None
    city: Optional[str] = None
    country: str = "Zimbabwe"
    notes: Optional[str] = None
    tags: Optional[str] = None

class MatterCreate(BaseModel):
    title: str
    client_id: int
    matter_type: str = "litigation"
    description: Optional[str] = None
    court_name: Optional[str] = None
    case_number: Optional[str] = None
    opposing_counsel: Optional[str] = None
    opposing_party: Optional[str] = None
    billing_type: str = "hourly"
    hourly_rate: Optional[float] = None
    is_urgent: bool = False
    is_pro_bono: bool = False
    tags: Optional[str] = None

class MatterNoteCreate(BaseModel):
    content: str
    note_type: str = "general"
    is_privileged: bool = True

class HearingCreate(BaseModel):
    hearing_type: str = "hearing"
    title: str
    date: str   # ISO format YYYY-MM-DD
    time: Optional[str] = None
    location: Optional[str] = None
    court_name: Optional[str] = None
    description: Optional[str] = None
    reminder_days: int = 3

class TaskCreate(BaseModel):
    title: str
    matter_id: Optional[int] = None
    description: Optional[str] = None
    priority: str = "medium"
    due_date: Optional[str] = None
    assigned_to: Optional[int] = None

class TimeEntryCreate(BaseModel):
    matter_id: int
    date: str
    hours: float
    rate: float
    description: str
    activity_code: Optional[str] = None
    is_billable: bool = True

class InvoiceCreate(BaseModel):
    client_id: int
    matter_id: Optional[int] = None
    due_date: Optional[str] = None
    notes: Optional[str] = None
    time_entry_ids: Optional[List[int]] = None
    disbursement_ids: Optional[List[int]] = None
    vat_rate: float = 0.15

class PaymentCreate(BaseModel):
    invoice_id: int
    amount: float
    payment_date: str
    payment_method: str = "transfer"
    reference: Optional[str] = None
    notes: Optional[str] = None

class TrustTransactionCreate(BaseModel):
    client_id: int
    matter_id: Optional[int] = None
    transaction_type: str
    amount: float
    date: str
    description: str
    reference: Optional[str] = None
    payment_method: Optional[str] = None

class DisbursementCreate(BaseModel):
    matter_id: int
    date: str
    description: str
    amount: float
    disbursement_type: Optional[str] = None
    receipt_ref: Optional[str] = None
    is_billable: bool = True

class ConflictCheckRequest(BaseModel):
    client_id: int
    opposing_party_name: Optional[str] = None
    opposing_counsel_name: Optional[str] = None

class ConflictRaiseRequest(BaseModel):
    client_id: int
    opposing_name: str
    opposing_counsel_name: Optional[str] = None
    reason: str
    risk_level: str = "medium"

class ConflictClearRequest(BaseModel):
    notes: Optional[str] = None

class ConflictDeclineRequest(BaseModel):
    notes: Optional[str] = None

class AIRequest(BaseModel):
    task_type: str
    prompt: str
    document_ids: Optional[List[int]] = None
    matter_id: Optional[int] = None
    context: Optional[str] = None
    max_tokens: int = 2500

class ContractDraftRequest(BaseModel):
    contract_type: str
    parties: str
    key_terms: str
    special_instructions: Optional[str] = None

class LetterDraftRequest(BaseModel):
    letter_type: str
    addressee: str
    subject: str
    key_points: str
    tone: str = "formal"

class ResearchRequest(BaseModel):
    query: str
    matter_id: Optional[int] = None
    jurisdiction: str = "Zimbabwe"
    research_type: str = "general"

class SettingsUpdate(BaseModel):
    firm_name: Optional[str] = None
    firm_address: Optional[str] = None
    firm_phone: Optional[str] = None
    firm_email: Optional[str] = None
    default_hourly_rate: Optional[float] = None
    vat_rate: Optional[float] = None
    invoice_prefix: Optional[str] = None
    trust_bank_name: Optional[str] = None
    trust_account_number: Optional[str] = None
    operating_bank_name: Optional[str] = None
    operating_account_number: Optional[str] = None


# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def _log_activity(db: Session, user_id: int, action: str, description: str,
                  resource_type: str = None, resource_id: int = None):
    db.add(Activity(user_id=user_id, action=action, description=description,
                    resource_type=resource_type, resource_id=resource_id))
    db.commit()


def _next_number(db: Session, prefix: str, model, field: str) -> str:
    """Generate sequential numbers like CLT-0001, MTR-2024-001, INV-001"""
    count = db.query(model).count() + 1
    if prefix == "MTR":
        return f"MTR-{datetime.now().year}-{count:04d}"
    return f"{prefix}-{count:04d}"


def _get_upload_dir(sub: str = "") -> Path:
    firm_id = os.getenv("FIRM_ID", "default")
    d = Path.home() / "IntelliLaw" / "firms" / firm_id / "documents" / sub
    d.mkdir(parents=True, exist_ok=True)
    return d


def _get_document_context(db: Session, document_ids: List[int]) -> str:
    if not document_ids:
        return ""
    docs = db.query(LegalDocument).filter(LegalDocument.id.in_(document_ids)).all()
    ctx = ""
    for doc in docs:
        ctx += f"\n{'='*60}\nDOCUMENT: {doc.original_filename}\n{'='*60}\n"
        if doc.extracted_text:
            ctx += doc.extracted_text[:3000]
        elif doc.extracted_text_path and Path(doc.extracted_text_path).exists():
            try:
                ctx += Path(doc.extracted_text_path).read_text(encoding="utf-8")[:3000]
            except Exception:
                pass
    return ctx[:6000]


# ─────────────────────────────────────────────────────────────
# HEALTH & SYSTEM
# ─────────────────────────────────────────────────────────────

@app.get("/api/health")
def health():
    return {"status": "ok", "app": "IntelliLaw", "version": "1.0.0"}


@app.get("/api/system/info")
def system_info(current_user: User = Depends(get_current_active_user)):
    try:
        from app.local_llm_service import get_ai_status
        ai = get_ai_status()
    except Exception:
        ai = {"available": False}
    return {
        "database": get_database_info(),
        "ai": ai,
        "demo": get_demo_status(),
    }


# ─────────────────────────────────────────────────────────────
# AUTH
# ─────────────────────────────────────────────────────────────

@app.post("/api/auth/login")
def login(req: LoginRequest, db: Session = Depends(get_db)):
    user = authenticate_user(db, req.username, req.password)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    token = create_access_token({"sub": user.username},
                                expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    _log_activity(db, user.id, "login", f"User {user.username} logged in")
    return {
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "id": user.id, "username": user.username,
            "full_name": user.full_name, "email": user.email,
            "user_role": user.user_role, "is_admin": user.is_admin,
            "title": user.title, "initials": user.initials,
        }
    }


@app.get("/api/auth/me")
def me(current_user: User = Depends(get_current_active_user)):
    return {
        "id": current_user.id, "username": current_user.username,
        "full_name": current_user.full_name, "email": current_user.email,
        "user_role": current_user.user_role, "is_admin": current_user.is_admin,
        "title": current_user.title, "initials": current_user.initials,
    }


@app.post("/api/auth/register")
def register(req: CreateUserRequest, db: Session = Depends(get_db),
             current_user: User = Depends(require_admin)):
    if db.query(User).filter(User.username == req.username).first():
        raise HTTPException(status_code=400, detail="Username already exists")
    user = User(
        username=req.username, email=req.email,
        hashed_password=get_password_hash(req.password),
        full_name=req.full_name, title=req.title, user_role=req.user_role,
        initials="".join(p[0].upper() for p in (req.full_name or req.username).split()[:2])
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return {"id": user.id, "username": user.username, "message": "User created"}


# ─────────────────────────────────────────────────────────────
# DASHBOARD
# ─────────────────────────────────────────────────────────────

@app.get("/api/dashboard/stats")
def dashboard_stats(db: Session = Depends(get_db),
                    current_user: User = Depends(get_current_active_user)):
    total_clients  = db.query(Client).filter(Client.is_active == True).count()
    active_matters = db.query(Matter).filter(Matter.status == "active").count()
    total_matters  = db.query(Matter).count()

    # Upcoming hearings (next 14 days)
    today = date.today()
    upcoming_hearings = db.query(Hearing).filter(
        Hearing.date >= today,
        Hearing.date <= today + timedelta(days=14),
        Hearing.is_completed == False
    ).order_by(Hearing.date).limit(10).all()

    # Overdue tasks
    overdue_tasks = db.query(Task).filter(
        Task.status.in_(["pending", "in_progress"]),
        Task.due_date < today
    ).count()

    # Outstanding invoices
    outstanding = db.query(func.sum(Invoice.balance_due)).filter(
        Invoice.status.in_(["sent", "partial", "overdue"])
    ).scalar() or 0.0

    # Recent AI sessions
    recent_ai = db.query(AISession).order_by(
        desc(AISession.created_at)).limit(5).all()

    # Matter types breakdown
    matter_types = db.query(
        Matter.matter_type, func.count(Matter.id)
    ).group_by(Matter.matter_type).all()

    return {
        "total_clients": total_clients,
        "active_matters": active_matters,
        "total_matters": total_matters,
        "overdue_tasks": overdue_tasks,
        "outstanding_invoices": round(outstanding, 2),
        "upcoming_hearings": [
            {
                "id": h.id, "title": h.title,
                "date": str(h.date), "time": h.time,
                "location": h.location, "hearing_type": h.hearing_type,
                "matter_id": h.matter_id,
            }
            for h in upcoming_hearings
        ],
        "recent_ai_sessions": [
            {
                "id": s.id, "type": s.session_type,
                "created_at": s.created_at.isoformat()
            }
            for s in recent_ai
        ],
        "matter_types_breakdown": [
            {"type": t, "count": c} for t, c in matter_types
        ],
    }


# ─────────────────────────────────────────────────────────────
# CLIENTS
# ─────────────────────────────────────────────────────────────

@app.get("/api/clients")
def list_clients(
    search: Optional[str] = None,
    client_type: Optional[str] = None,
    skip: int = 0, limit: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    q = db.query(Client).filter(Client.is_active == True)
    if search:
        q = q.filter(or_(
            Client.first_name.ilike(f"%{search}%"),
            Client.last_name.ilike(f"%{search}%"),
            Client.company_name.ilike(f"%{search}%"),
            Client.email.ilike(f"%{search}%"),
            Client.client_number.ilike(f"%{search}%"),
        ))
    if client_type:
        q = q.filter(Client.client_type == client_type)
    total = q.count()
    clients = q.order_by(Client.last_name, Client.company_name).offset(skip).limit(limit).all()
    return {
        "total": total,
        "items": [_client_out(c, db) for c in clients]
    }


@app.post("/api/clients", status_code=201)
def create_client(req: ClientCreate, db: Session = Depends(get_db),
                  current_user: User = Depends(get_current_active_user)):
    num = _next_number(db, "CLT", Client, "client_number")
    client = Client(**req.dict(), client_number=num, created_by=current_user.id)
    db.add(client)
    db.commit()
    db.refresh(client)
    _log_activity(db, current_user.id, "create_client",
                  f"Created client {client.display_name}", "client", client.id)
    return _client_out(client, db)


@app.get("/api/clients/{client_id}")
def get_client(client_id: int, db: Session = Depends(get_db),
               current_user: User = Depends(get_current_active_user)):
    client = db.query(Client).filter(Client.id == client_id).first()
    if not client:
        raise HTTPException(404, "Client not found")
    return _client_out(client, db, detailed=True)


@app.put("/api/clients/{client_id}")
def update_client(client_id: int, req: ClientCreate,
                  db: Session = Depends(get_db),
                  current_user: User = Depends(get_current_active_user)):
    client = db.query(Client).filter(Client.id == client_id).first()
    if not client:
        raise HTTPException(404, "Client not found")
    for k, v in req.dict(exclude_unset=True).items():
        setattr(client, k, v)
    client.updated_at = datetime.utcnow()
    db.commit()
    return _client_out(client, db)


def _client_out(c: Client, db: Session, detailed: bool = False) -> dict:
    matters_count = db.query(Matter).filter(Matter.client_id == c.id).count()
    d = {
        "id": c.id, "client_number": c.client_number,
        "client_type": c.client_type,
        "display_name": c.display_name,
        "first_name": c.first_name, "last_name": c.last_name,
        "company_name": c.company_name,
        "email": c.email, "phone": c.phone,
        "city": c.city, "country": c.country,
        "is_active": c.is_active, "risk_rating": c.risk_rating,
        "kyc_verified": c.kyc_verified,
        "tags": c.tags, "notes": c.notes,
        "matters_count": matters_count,
        "created_at": c.created_at.isoformat() if c.created_at else None,
    }
    if detailed:
        d["address_line1"] = c.address_line1
        d["address_line2"] = c.address_line2
        d["id_number"] = c.id_number
        d["registration_no"] = c.registration_no
        d["contact_person"] = c.contact_person
    return d


# ─────────────────────────────────────────────────────────────
# MATTERS
# ─────────────────────────────────────────────────────────────

@app.get("/api/matters")
def list_matters(
    search: Optional[str] = None,
    status: Optional[str] = None,
    client_id: Optional[int] = None,
    matter_type: Optional[str] = None,
    skip: int = 0, limit: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    q = db.query(Matter)
    if search:
        q = q.filter(or_(
            Matter.title.ilike(f"%{search}%"),
            Matter.matter_number.ilike(f"%{search}%"),
            Matter.case_number.ilike(f"%{search}%"),
        ))
    if status:
        q = q.filter(Matter.status == status)
    if client_id:
        q = q.filter(Matter.client_id == client_id)
    if matter_type:
        q = q.filter(Matter.matter_type == matter_type)
    total = q.count()
    matters = q.order_by(desc(Matter.updated_at)).offset(skip).limit(limit).all()
    return {"total": total, "items": [_matter_out(m, db) for m in matters]}


@app.post("/api/matters", status_code=201)
def create_matter(req: MatterCreate, db: Session = Depends(get_db),
                  current_user: User = Depends(get_current_active_user)):
    # Check for uncleared conflicts
    uncleared = db.query(Conflict).filter(
        (Conflict.client_id == req.client_id) &
        (Conflict.status.in_(["raised", "under_review"]))
    ).all()

    if uncleared:
        raise HTTPException(
            status_code=409,
            detail=f"Cannot create matter: {len(uncleared)} uncleared conflict(s) exist. "
                   f"Resolve in Conflict Check module."
        )

    num = _next_number(db, "MTR", Matter, "matter_number")
    matter = Matter(
        **req.dict(),
        matter_number=num,
        responsible_attorney_id=current_user.id,
        created_by=current_user.id,
    )
    db.add(matter)
    db.commit()
    db.refresh(matter)
    _log_activity(db, current_user.id, "create_matter",
                  f"Opened matter {matter.matter_number}: {matter.title}", "matter", matter.id)
    return _matter_out(matter, db)


@app.get("/api/matters/{matter_id}")
def get_matter(matter_id: int, db: Session = Depends(get_db),
               current_user: User = Depends(get_current_active_user)):
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(404, "Matter not found")
    return _matter_out(matter, db, detailed=True)


@app.put("/api/matters/{matter_id}")
def update_matter(matter_id: int, req: MatterCreate,
                  db: Session = Depends(get_db),
                  current_user: User = Depends(get_current_active_user)):
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(404, "Matter not found")
    for k, v in req.dict(exclude_unset=True).items():
        setattr(matter, k, v)
    matter.updated_at = datetime.utcnow()
    db.commit()
    return _matter_out(matter, db)


def _matter_out(m: Matter, db: Session, detailed: bool = False) -> dict:
    client = db.query(Client).filter(Client.id == m.client_id).first()
    attorney = db.query(User).filter(User.id == m.responsible_attorney_id).first() if m.responsible_attorney_id else None
    docs_count = db.query(LegalDocument).filter(LegalDocument.matter_id == m.id).count()
    d = {
        "id": m.id, "matter_number": m.matter_number,
        "title": m.title, "matter_type": m.matter_type, "status": m.status,
        "client_id": m.client_id,
        "client_name": client.display_name if client else None,
        "responsible_attorney": attorney.full_name if attorney else None,
        "court_name": m.court_name, "case_number": m.case_number,
        "opposing_party": m.opposing_party,
        "billing_type": m.billing_type, "hourly_rate": m.hourly_rate,
        "is_urgent": m.is_urgent, "is_pro_bono": m.is_pro_bono,
        "tags": m.tags, "documents_count": docs_count,
        "opened_date": str(m.opened_date) if m.opened_date else None,
        "next_hearing_date": str(m.next_hearing_date) if m.next_hearing_date else None,
        "updated_at": m.updated_at.isoformat() if m.updated_at else None,
    }
    if detailed:
        d["description"] = m.description
        d["internal_notes"] = m.internal_notes
        d["opposing_counsel"] = m.opposing_counsel
        d["judge_name"] = m.judge_name
        d["division"] = m.division
        d["estimated_value"] = m.estimated_value
        d["retainer_amount"] = m.retainer_amount
        d["fixed_fee"] = m.fixed_fee
    return d


# ─────────────────────────────────────────────────────────────
# CONFLICTS
# ─────────────────────────────────────────────────────────────

@app.post("/api/conflicts/check", status_code=200)
def check_conflicts(req: ConflictCheckRequest,
                   db: Session = Depends(get_db),
                   current_user: User = Depends(get_current_active_user)):
    """Search for potential conflicts before matter creation"""
    result = search_potential_conflicts(
        db,
        client_id=req.client_id,
        opposing_party_name=req.opposing_party_name or "",
        opposing_counsel_name=req.opposing_counsel_name or ""
    )
    return result


@app.get("/api/conflicts")
def list_conflicts(
    status: Optional[str] = None,
    client_id: Optional[int] = None,
    skip: int = 0, limit: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """List conflicts with optional filters"""
    q = db.query(Conflict)
    if status:
        q = q.filter(Conflict.status == status)
    if client_id:
        q = q.filter(Conflict.client_id == client_id)
    total = q.count()
    conflicts = q.order_by(desc(Conflict.raised_at)).offset(skip).limit(limit).all()

    items = []
    for c in conflicts:
        items.append({
            "id": c.id,
            "status": c.status,
            "client_id": c.client_id,
            "opposing_name": c.opposing_name,
            "opposing_counsel_name": c.opposing_counsel_name,
            "risk_level": c.risk_level,
            "reason": c.reason,
            "matter_id": c.matter_id,
            "raised_by_id": c.raised_by_id,
            "raised_at": c.raised_at.isoformat() if c.raised_at else None,
            "reviewed_by_id": c.reviewed_by_id,
            "reviewed_at": c.reviewed_at.isoformat() if c.reviewed_at else None,
            "notes": c.notes,
            "created_at": c.created_at.isoformat() if c.created_at else None,
        })
    return {"total": total, "items": items}


@app.post("/api/conflicts", status_code=201)
def raise_conflict(req: ConflictRaiseRequest,
                   db: Session = Depends(get_db),
                   current_user: User = Depends(get_current_active_user)):
    """Raise a new conflict of interest"""
    conflict = Conflict(
        client_id=req.client_id,
        opposing_name=req.opposing_name,
        opposing_counsel_name=req.opposing_counsel_name,
        reason=req.reason,
        risk_level=req.risk_level,
        status="raised",
        raised_by_id=current_user.id,
    )
    db.add(conflict)
    db.commit()
    db.refresh(conflict)
    _log_activity(db, current_user.id, "raise_conflict",
                  f"Raised conflict: {req.opposing_name}", "conflict", conflict.id)
    return {
        "id": conflict.id,
        "status": conflict.status,
        "client_id": conflict.client_id,
        "opposing_name": conflict.opposing_name,
        "opposing_counsel_name": conflict.opposing_counsel_name,
        "risk_level": conflict.risk_level,
        "reason": conflict.reason,
        "raised_by_id": conflict.raised_by_id,
        "raised_at": conflict.raised_at.isoformat() if conflict.raised_at else None,
    }


@app.put("/api/conflicts/{conflict_id}/clear")
def clear_conflict(conflict_id: int, req: ConflictClearRequest,
                   db: Session = Depends(get_db),
                   current_user: User = Depends(get_current_active_user)):
    """Clear a conflict (attorney/admin only)"""
    conflict = db.query(Conflict).filter(Conflict.id == conflict_id).first()
    if not conflict:
        raise HTTPException(404, "Conflict not found")

    conflict.status = "cleared"
    conflict.reviewed_by_id = current_user.id
    conflict.reviewed_at = datetime.utcnow()
    conflict.notes = req.notes
    db.commit()
    db.refresh(conflict)
    _log_activity(db, current_user.id, "clear_conflict",
                  f"Cleared conflict: {conflict.opposing_name}", "conflict", conflict.id)
    return {
        "id": conflict.id,
        "status": conflict.status,
        "reviewed_by_id": conflict.reviewed_by_id,
        "reviewed_at": conflict.reviewed_at.isoformat() if conflict.reviewed_at else None,
        "notes": conflict.notes,
    }


@app.put("/api/conflicts/{conflict_id}/decline")
def decline_conflict(conflict_id: int, req: ConflictDeclineRequest,
                    db: Session = Depends(get_db),
                    current_user: User = Depends(get_current_active_user)):
    """Decline a conflict (attorney/admin only)"""
    conflict = db.query(Conflict).filter(Conflict.id == conflict_id).first()
    if not conflict:
        raise HTTPException(404, "Conflict not found")

    conflict.status = "declined"
    conflict.reviewed_by_id = current_user.id
    conflict.reviewed_at = datetime.utcnow()
    conflict.notes = req.notes
    db.commit()
    db.refresh(conflict)
    _log_activity(db, current_user.id, "decline_conflict",
                  f"Declined conflict: {conflict.opposing_name}", "conflict", conflict.id)
    return {
        "id": conflict.id,
        "status": conflict.status,
        "reviewed_by_id": conflict.reviewed_by_id,
        "reviewed_at": conflict.reviewed_at.isoformat() if conflict.reviewed_at else None,
        "notes": conflict.notes,
    }


# ─────────────────────────────────────────────────────────────
# MATTER — NOTES
# ─────────────────────────────────────────────────────────────

@app.get("/api/matters/{matter_id}/notes")
def get_notes(matter_id: int, db: Session = Depends(get_db),
              current_user: User = Depends(get_current_active_user)):
    notes = db.query(MatterNote).filter(MatterNote.matter_id == matter_id).order_by(
        desc(MatterNote.created_at)).all()
    return [_note_out(n, db) for n in notes]


@app.post("/api/matters/{matter_id}/notes", status_code=201)
def add_note(matter_id: int, req: MatterNoteCreate,
             db: Session = Depends(get_db),
             current_user: User = Depends(get_current_active_user)):
    note = MatterNote(matter_id=matter_id, author_id=current_user.id, **req.dict())
    db.add(note)
    db.commit()
    db.refresh(note)
    return _note_out(note, db)


def _note_out(n: MatterNote, db: Session) -> dict:
    author = db.query(User).filter(User.id == n.author_id).first()
    return {
        "id": n.id, "matter_id": n.matter_id,
        "content": n.content, "note_type": n.note_type,
        "is_privileged": n.is_privileged,
        "author": author.full_name if author else "Unknown",
        "created_at": n.created_at.isoformat(),
    }


# ─────────────────────────────────────────────────────────────
# HEARINGS
# ─────────────────────────────────────────────────────────────

@app.get("/api/matters/{matter_id}/hearings")
def get_hearings(matter_id: int, db: Session = Depends(get_db),
                 current_user: User = Depends(get_current_active_user)):
    hearings = db.query(Hearing).filter(Hearing.matter_id == matter_id).order_by(
        Hearing.date).all()
    return [_hearing_out(h) for h in hearings]


@app.post("/api/matters/{matter_id}/hearings", status_code=201)
def add_hearing(matter_id: int, req: HearingCreate,
                db: Session = Depends(get_db),
                current_user: User = Depends(get_current_active_user)):
    h = Hearing(
        matter_id=matter_id, created_by=current_user.id,
        **{k: v for k, v in req.dict().items() if k != "date"},
        date=date.fromisoformat(req.date),
    )
    db.add(h)
    # Update matter's next hearing
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if matter:
        matter.next_hearing_date = date.fromisoformat(req.date)
    db.commit()
    db.refresh(h)
    return _hearing_out(h)


@app.get("/api/calendar/hearings")
def all_hearings(
    from_date: Optional[str] = None, to_date: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    q = db.query(Hearing).filter(Hearing.is_completed == False)
    if from_date:
        q = q.filter(Hearing.date >= date.fromisoformat(from_date))
    if to_date:
        q = q.filter(Hearing.date <= date.fromisoformat(to_date))
    hearings = q.order_by(Hearing.date).limit(100).all()
    return [_hearing_out(h) for h in hearings]


def _hearing_out(h: Hearing) -> dict:
    return {
        "id": h.id, "matter_id": h.matter_id,
        "hearing_type": h.hearing_type, "title": h.title,
        "date": str(h.date), "time": h.time,
        "duration_minutes": h.duration_minutes,
        "location": h.location, "court_name": h.court_name,
        "judge_name": h.judge_name,
        "description": h.description, "outcome": h.outcome,
        "is_completed": h.is_completed,
        "reminder_days": h.reminder_days,
    }


# ─────────────────────────────────────────────────────────────
# TASKS
# ─────────────────────────────────────────────────────────────

@app.get("/api/tasks")
def list_tasks(
    status: Optional[str] = None, priority: Optional[str] = None,
    matter_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    q = db.query(Task)
    if status:
        q = q.filter(Task.status == status)
    if priority:
        q = q.filter(Task.priority == priority)
    if matter_id:
        q = q.filter(Task.matter_id == matter_id)
    tasks = q.order_by(Task.due_date).limit(100).all()
    return [_task_out(t, db) for t in tasks]


@app.post("/api/tasks", status_code=201)
def create_task(req: TaskCreate, db: Session = Depends(get_db),
                current_user: User = Depends(get_current_active_user)):
    t = Task(
        created_by=current_user.id,
        **{k: v for k, v in req.dict().items() if k != "due_date"},
        due_date=date.fromisoformat(req.due_date) if req.due_date else None,
    )
    db.add(t)
    db.commit()
    db.refresh(t)
    return _task_out(t, db)


@app.put("/api/tasks/{task_id}/complete")
def complete_task(task_id: int, db: Session = Depends(get_db),
                  current_user: User = Depends(get_current_active_user)):
    t = db.query(Task).filter(Task.id == task_id).first()
    if not t:
        raise HTTPException(404, "Task not found")
    t.status = "completed"
    t.completed_date = date.today()
    db.commit()
    return _task_out(t, db)


def _task_out(t: Task, db: Session) -> dict:
    assignee = db.query(User).filter(User.id == t.assigned_to).first() if t.assigned_to else None
    return {
        "id": t.id, "title": t.title, "description": t.description,
        "matter_id": t.matter_id, "priority": t.priority, "status": t.status,
        "due_date": str(t.due_date) if t.due_date else None,
        "completed_date": str(t.completed_date) if t.completed_date else None,
        "assignee": assignee.full_name if assignee else None,
        "created_at": t.created_at.isoformat(),
    }


# ─────────────────────────────────────────────────────────────
# DOCUMENTS
# ─────────────────────────────────────────────────────────────

@app.post("/api/documents/upload", status_code=201)
async def upload_document(
    matter_id: Optional[int] = Form(None),
    client_id: Optional[int] = Form(None),
    doc_category: str = Form("general"),
    description: Optional[str] = Form(None),
    is_privileged: bool = Form(False),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    ext = Path(file.filename).suffix.lower()
    unique_name = f"{uuid.uuid4().hex}{ext}"
    upload_dir = _get_upload_dir(str(matter_id) if matter_id else "general")
    file_path = upload_dir / unique_name

    content = await file.read()
    file_path.write_bytes(content)

    # Extract text
    extracted_text = None
    try:
        if ext == ".pdf":
            import pypdf, io
            reader = pypdf.PdfReader(io.BytesIO(content))
            extracted_text = "\n".join(p.extract_text() or "" for p in reader.pages)
        elif ext in (".txt", ".md"):
            extracted_text = content.decode("utf-8", errors="replace")
        elif ext in (".docx",):
            import docx, io
            doc_obj = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join(p.text for p in doc_obj.paragraphs)
    except Exception:
        pass

    doc = LegalDocument(
        matter_id=matter_id, client_id=client_id,
        filename=unique_name, original_filename=file.filename,
        file_path=str(file_path), file_type=ext.lstrip("."),
        file_size=len(content),
        extracted_text=extracted_text,
        doc_category=doc_category, description=description,
        is_privileged=is_privileged,
        uploaded_by=current_user.id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    _log_activity(db, current_user.id, "upload_document",
                  f"Uploaded {file.filename}", "document", doc.id)
    return _doc_out(doc)


@app.get("/api/documents")
def list_documents(
    matter_id: Optional[int] = None,
    client_id: Optional[int] = None,
    doc_category: Optional[str] = None,
    search: Optional[str] = None,
    skip: int = 0, limit: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    q = db.query(LegalDocument)
    if matter_id:
        q = q.filter(LegalDocument.matter_id == matter_id)
    if client_id:
        q = q.filter(LegalDocument.client_id == client_id)
    if doc_category:
        q = q.filter(LegalDocument.doc_category == doc_category)
    if search:
        q = q.filter(or_(
            LegalDocument.original_filename.ilike(f"%{search}%"),
            LegalDocument.description.ilike(f"%{search}%"),
        ))
    total = q.count()
    docs = q.order_by(desc(LegalDocument.upload_date)).offset(skip).limit(limit).all()
    return {"total": total, "items": [_doc_out(d) for d in docs]}


@app.get("/api/documents/{doc_id}/download")
def download_document(doc_id: int, db: Session = Depends(get_db),
                      current_user: User = Depends(get_current_active_user)):
    doc = db.query(LegalDocument).filter(LegalDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Document not found")
    doc.usage_count = (doc.usage_count or 0) + 1
    doc.last_used = datetime.utcnow()
    db.commit()
    return FileResponse(doc.file_path, filename=doc.original_filename)


def _doc_out(d: LegalDocument) -> dict:
    return {
        "id": d.id, "original_filename": d.original_filename,
        "file_type": d.file_type, "file_size": d.file_size,
        "doc_category": d.doc_category, "description": d.description,
        "matter_id": d.matter_id, "client_id": d.client_id,
        "is_privileged": d.is_privileged,
        "has_text": bool(d.extracted_text),
        "usage_count": d.usage_count,
        "upload_date": d.upload_date.isoformat() if d.upload_date else None,
    }


# ─────────────────────────────────────────────────────────────
# TIME ENTRIES
# ─────────────────────────────────────────────────────────────

@app.get("/api/billing/time-entries")
def list_time_entries(
    matter_id: Optional[int] = None, is_billed: Optional[bool] = None,
    skip: int = 0, limit: int = 100,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    q = db.query(TimeEntry)
    if matter_id:
        q = q.filter(TimeEntry.matter_id == matter_id)
    if is_billed is not None:
        q = q.filter(TimeEntry.is_billed == is_billed)
    total = q.count()
    entries = q.order_by(desc(TimeEntry.date)).offset(skip).limit(limit).all()
    return {"total": total, "items": [_time_entry_out(e, db) for e in entries]}


@app.post("/api/billing/time-entries", status_code=201)
def create_time_entry(req: TimeEntryCreate, db: Session = Depends(get_db),
                      current_user: User = Depends(get_current_active_user)):
    entry = TimeEntry(
        attorney_id=current_user.id,
        amount=req.hours * req.rate,
        **{k: v for k, v in req.dict().items() if k != "date"},
        date=date.fromisoformat(req.date),
    )
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return _time_entry_out(entry, db)


def _time_entry_out(e: TimeEntry, db: Session) -> dict:
    attorney = db.query(User).filter(User.id == e.attorney_id).first()
    matter = db.query(Matter).filter(Matter.id == e.matter_id).first()
    return {
        "id": e.id, "matter_id": e.matter_id,
        "matter_title": matter.title if matter else None,
        "attorney": attorney.full_name if attorney else None,
        "date": str(e.date), "hours": e.hours,
        "rate": e.rate, "amount": e.amount,
        "description": e.description,
        "activity_code": e.activity_code,
        "is_billable": e.is_billable, "is_billed": e.is_billed,
    }


# ─────────────────────────────────────────────────────────────
# INVOICES
# ─────────────────────────────────────────────────────────────

@app.get("/api/billing/invoices")
def list_invoices(
    client_id: Optional[int] = None, status: Optional[str] = None,
    skip: int = 0, limit: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    q = db.query(Invoice)
    if client_id:
        q = q.filter(Invoice.client_id == client_id)
    if status:
        q = q.filter(Invoice.status == status)
    total = q.count()
    invoices = q.order_by(desc(Invoice.issue_date)).offset(skip).limit(limit).all()
    return {"total": total, "items": [_invoice_out(i, db) for i in invoices]}


@app.post("/api/billing/invoices", status_code=201)
def create_invoice(req: InvoiceCreate, db: Session = Depends(get_db),
                   current_user: User = Depends(get_current_active_user)):
    settings = db.query(SystemSettings).first()
    prefix = settings.invoice_prefix if settings else "INV"
    num = _next_number(db, prefix, Invoice, "invoice_number")

    subtotal = 0.0
    line_items = []

    # Add time entries
    if req.time_entry_ids:
        entries = db.query(TimeEntry).filter(TimeEntry.id.in_(req.time_entry_ids)).all()
        for e in entries:
            subtotal += e.amount
            line_items.append(InvoiceLineItem(
                description=f"{e.description} — {e.hours}h @ {e.rate}/hr",
                quantity=e.hours, unit_price=e.rate, amount=e.amount, item_type="fee"
            ))

    # Add disbursements
    if req.disbursement_ids:
        disbs = db.query(Disbursement).filter(Disbursement.id.in_(req.disbursement_ids)).all()
        for d in disbs:
            subtotal += d.amount
            line_items.append(InvoiceLineItem(
                description=d.description,
                quantity=1, unit_price=d.amount, amount=d.amount, item_type="disbursement"
            ))

    vat_amount = subtotal * req.vat_rate
    total = subtotal + vat_amount

    invoice = Invoice(
        invoice_number=num, client_id=req.client_id,
        matter_id=req.matter_id, status="draft",
        due_date=date.fromisoformat(req.due_date) if req.due_date else date.today() + timedelta(days=30),
        notes=req.notes, vat_rate=req.vat_rate,
        subtotal=subtotal, vat_amount=vat_amount,
        total=total, balance_due=total,
        created_by=current_user.id,
    )
    db.add(invoice)
    db.flush()

    for li in line_items:
        li.invoice_id = invoice.id
        db.add(li)

    # Mark time entries as billed
    if req.time_entry_ids:
        db.query(TimeEntry).filter(TimeEntry.id.in_(req.time_entry_ids)).update(
            {"is_billed": True, "invoice_id": invoice.id}
        )

    db.commit()
    db.refresh(invoice)
    _log_activity(db, current_user.id, "create_invoice",
                  f"Created invoice {invoice.invoice_number}", "invoice", invoice.id)
    return _invoice_out(invoice, db)


@app.post("/api/billing/invoices/{invoice_id}/send")
def send_invoice(invoice_id: int, db: Session = Depends(get_db),
                 current_user: User = Depends(get_current_active_user)):
    inv = db.query(Invoice).filter(Invoice.id == invoice_id).first()
    if not inv:
        raise HTTPException(404, "Invoice not found")
    inv.status = "sent"
    db.commit()
    return {"message": "Invoice marked as sent", "invoice_number": inv.invoice_number}


def _invoice_out(i: Invoice, db: Session) -> dict:
    client = db.query(Client).filter(Client.id == i.client_id).first()
    matter = db.query(Matter).filter(Matter.id == i.matter_id).first() if i.matter_id else None
    return {
        "id": i.id, "invoice_number": i.invoice_number,
        "status": i.status,
        "client_id": i.client_id,
        "client_name": client.display_name if client else None,
        "matter_id": i.matter_id,
        "matter_title": matter.title if matter else None,
        "issue_date": str(i.issue_date), "due_date": str(i.due_date),
        "subtotal": i.subtotal, "vat_amount": i.vat_amount,
        "total": i.total, "amount_paid": i.amount_paid, "balance_due": i.balance_due,
        "currency": i.currency, "notes": i.notes,
    }


# ─────────────────────────────────────────────────────────────
# PAYMENTS
# ─────────────────────────────────────────────────────────────

@app.post("/api/billing/payments", status_code=201)
def record_payment(req: PaymentCreate, db: Session = Depends(get_db),
                   current_user: User = Depends(get_current_active_user)):
    inv = db.query(Invoice).filter(Invoice.id == req.invoice_id).first()
    if not inv:
        raise HTTPException(404, "Invoice not found")
    payment = Payment(
        client_id=inv.client_id, created_by=current_user.id,
        **{k: v for k, v in req.dict().items() if k != "payment_date"},
        payment_date=date.fromisoformat(req.payment_date),
    )
    db.add(payment)
    inv.amount_paid = (inv.amount_paid or 0) + req.amount
    inv.balance_due = inv.total - inv.amount_paid
    inv.status = "paid" if inv.balance_due <= 0 else "partial"
    db.commit()
    return {"message": "Payment recorded", "balance_due": inv.balance_due}


# ─────────────────────────────────────────────────────────────
# TRUST ACCOUNTS
# ─────────────────────────────────────────────────────────────

@app.get("/api/billing/trust")
def list_trust(db: Session = Depends(get_db),
               current_user: User = Depends(get_current_active_user)):
    accounts = db.query(TrustAccount).all()
    return [_trust_out(a, db) for a in accounts]


@app.post("/api/billing/trust/transaction", status_code=201)
def trust_transaction(req: TrustTransactionCreate, db: Session = Depends(get_db),
                      current_user: User = Depends(get_current_active_user)):
    account = db.query(TrustAccount).filter(
        TrustAccount.client_id == req.client_id).first()
    if not account:
        account = TrustAccount(client_id=req.client_id, matter_id=req.matter_id, balance=0.0)
        db.add(account)
        db.flush()

    if req.transaction_type == "disbursement" and account.balance < req.amount:
        raise HTTPException(400, "Insufficient trust balance")

    delta = req.amount if req.transaction_type == "receipt" else -req.amount
    account.balance += delta

    txn = TrustTransaction(
        account_id=account.id,
        transaction_type=req.transaction_type,
        amount=req.amount,
        running_balance=account.balance,
        description=req.description,
        reference=req.reference,
        payment_method=req.payment_method,
        created_by=current_user.id,
        date=date.fromisoformat(req.date),
    )
    db.add(txn)
    db.commit()
    _log_activity(db, current_user.id, "trust_transaction",
                  f"Trust {req.transaction_type}: {req.amount} for client {req.client_id}")
    return {"message": "Transaction recorded", "new_balance": account.balance}


def _trust_out(a: TrustAccount, db: Session) -> dict:
    client = db.query(Client).filter(Client.id == a.client_id).first()
    return {
        "id": a.id, "client_id": a.client_id,
        "client_name": client.display_name if client else None,
        "balance": a.balance, "currency": a.currency,
        "opened_date": str(a.opened_date),
    }


# ─────────────────────────────────────────────────────────────
# DISBURSEMENTS
# ─────────────────────────────────────────────────────────────

@app.post("/api/billing/disbursements", status_code=201)
def create_disbursement(req: DisbursementCreate, db: Session = Depends(get_db),
                        current_user: User = Depends(get_current_active_user)):
    d = Disbursement(
        created_by=current_user.id,
        **{k: v for k, v in req.dict().items() if k != "date"},
        date=date.fromisoformat(req.date),
    )
    db.add(d)
    db.commit()
    db.refresh(d)
    return {"id": d.id, "amount": d.amount, "description": d.description}


# ─────────────────────────────────────────────────────────────
# BILLING REPORTS
# ─────────────────────────────────────────────────────────────

@app.get("/api/billing/reports/summary")
def billing_summary(db: Session = Depends(get_db),
                    current_user: User = Depends(get_current_active_user)):
    total_invoiced = db.query(func.sum(Invoice.total)).filter(
        Invoice.status != "cancelled").scalar() or 0
    total_received = db.query(func.sum(Invoice.amount_paid)).scalar() or 0
    outstanding     = db.query(func.sum(Invoice.balance_due)).filter(
        Invoice.status.in_(["sent", "partial", "overdue"])).scalar() or 0
    trust_total     = db.query(func.sum(TrustAccount.balance)).scalar() or 0

    unbilled_hours = db.query(func.sum(TimeEntry.hours)).filter(
        TimeEntry.is_billable == True,
        TimeEntry.is_billed == False
    ).scalar() or 0

    return {
        "total_invoiced": round(total_invoiced, 2),
        "total_received": round(total_received, 2),
        "outstanding": round(outstanding, 2),
        "trust_total": round(trust_total, 2),
        "unbilled_hours": round(unbilled_hours, 2),
    }


# ─────────────────────────────────────────────────────────────
# AI LEGAL ASSISTANT
# ─────────────────────────────────────────────────────────────

@app.post("/api/ai/generate")
async def ai_generate(req: AIRequest, db: Session = Depends(get_db),
                      current_user: User = Depends(get_current_active_user)):
    import time
    system_prompt = get_legal_prompt(req.task_type)
    doc_context   = _get_document_context(db, req.document_ids or [])

    full_prompt = req.prompt
    if doc_context:
        full_prompt = f"DOCUMENTS:\n{doc_context}\n\nINSTRUCTION:\n{req.prompt}"
    if req.context:
        full_prompt = f"CONTEXT:\n{req.context}\n\n{full_prompt}"

    start = time.time()
    try:
        from app.local_llm_service import generate_text as local_gen
        output = smart_generate(
            prompt=full_prompt,
            system_prompt=system_prompt,
            max_tokens=req.max_tokens,
            local_generate_func=lambda p, s, m: local_gen(p, max_tokens=m, system_message=s),
        )
    except Exception as e:
        raise HTTPException(500, f"AI generation failed: {e}")

    elapsed = time.time() - start
    session = AISession(
        session_type=req.task_type,
        input_prompt=full_prompt[:1000],
        output_content=output,
        document_ids=",".join(map(str, req.document_ids)) if req.document_ids else None,
        matter_id=req.matter_id,
        generation_time=elapsed,
        user_id=current_user.id,
    )
    db.add(session)
    db.commit()
    _log_activity(db, current_user.id, f"ai_{req.task_type}",
                  f"AI task: {req.task_type}", "ai_session", session.id)
    return {"content": output, "session_id": session.id, "generation_time": round(elapsed, 1)}


@app.post("/api/ai/draft-contract")
async def draft_contract(req: ContractDraftRequest, db: Session = Depends(get_db),
                         current_user: User = Depends(get_current_active_user)):
    prompt = build_draft_contract_prompt(
        req.contract_type, req.parties, req.key_terms, req.special_instructions)
    system = get_legal_prompt("draft_contract")
    try:
        from app.local_llm_service import generate_text as local_gen
        output = smart_generate(
            prompt=prompt, system_prompt=system, max_tokens=3000,
            local_generate_func=lambda p, s, m: local_gen(p, max_tokens=m, system_message=s),
        )
    except Exception as e:
        raise HTTPException(500, str(e))
    return {"content": output}


@app.post("/api/ai/draft-letter")
async def draft_letter(req: LetterDraftRequest, db: Session = Depends(get_db),
                       current_user: User = Depends(get_current_active_user)):
    prompt = build_draft_letter_prompt(
        req.letter_type, req.addressee, req.subject, req.key_points, req.tone)
    system = get_legal_prompt("draft_letter")
    try:
        from app.local_llm_service import generate_text as local_gen
        output = smart_generate(
            prompt=prompt, system_prompt=system, max_tokens=2000,
            local_generate_func=lambda p, s, m: local_gen(p, max_tokens=m, system_message=s),
        )
    except Exception as e:
        raise HTTPException(500, str(e))
    return {"content": output}


@app.post("/api/ai/summarize-document/{doc_id}")
async def summarize_document(doc_id: int, focus: Optional[str] = None,
                              db: Session = Depends(get_db),
                              current_user: User = Depends(get_current_active_user)):
    doc = db.query(LegalDocument).filter(LegalDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Document not found")
    text = doc.extracted_text or ""
    if not text:
        raise HTTPException(400, "Document has no extracted text")
    prompt = build_summarize_prompt(text, focus or "")
    system = get_legal_prompt("summarize")
    try:
        from app.local_llm_service import generate_text as local_gen
        output = smart_generate(
            prompt=prompt, system_prompt=system, max_tokens=2000,
            local_generate_func=lambda p, s, m: local_gen(p, max_tokens=m, system_message=s),
        )
    except Exception as e:
        raise HTTPException(500, str(e))
    return {"summary": output, "doc_id": doc_id}


@app.get("/api/ai/status")
def ai_status(current_user: User = Depends(get_current_active_user)):
    try:
        from app.local_llm_service import get_ai_status
        return {**get_ai_status(), "demo": get_demo_status()}
    except Exception:
        return {"available": False, "demo": get_demo_status()}


@app.get("/api/ai/sessions")
def ai_sessions(matter_id: Optional[int] = None,
                skip: int = 0, limit: int = 30,
                db: Session = Depends(get_db),
                current_user: User = Depends(get_current_active_user)):
    q = db.query(AISession).filter(AISession.user_id == current_user.id)
    if matter_id:
        q = q.filter(AISession.matter_id == matter_id)
    sessions = q.order_by(desc(AISession.created_at)).offset(skip).limit(limit).all()
    return [
        {
            "id": s.id, "session_type": s.session_type,
            "input_prompt": s.input_prompt[:200] if s.input_prompt else None,
            "output_content": s.output_content,
            "matter_id": s.matter_id,
            "generation_time": s.generation_time,
            "created_at": s.created_at.isoformat(),
        }
        for s in sessions
    ]


# ─────────────────────────────────────────────────────────────
# RESEARCH
# ─────────────────────────────────────────────────────────────

@app.post("/api/research/search")
async def research_search(req: ResearchRequest, db: Session = Depends(get_db),
                          current_user: User = Depends(get_current_active_user)):
    prompt = build_research_prompt(req.query, req.jurisdiction, req.research_type)
    system = get_legal_prompt("research")
    try:
        from app.local_llm_service import generate_text as local_gen
        output = smart_generate(
            prompt=prompt, system_prompt=system, max_tokens=3000,
            local_generate_func=lambda p, s, m: local_gen(p, max_tokens=m, system_message=s),
        )
    except Exception as e:
        raise HTTPException(500, str(e))

    session = ResearchSession(
        matter_id=req.matter_id, title=req.query[:200],
        query=req.query, ai_summary=output,
        research_type=req.research_type, jurisdiction=req.jurisdiction,
        created_by=current_user.id,
    )
    db.add(session)
    db.commit()
    return {"results": output, "session_id": session.id}


@app.get("/api/research/sessions")
def research_sessions(skip: int = 0, limit: int = 20,
                      db: Session = Depends(get_db),
                      current_user: User = Depends(get_current_active_user)):
    sessions = db.query(ResearchSession).filter(
        ResearchSession.created_by == current_user.id
    ).order_by(desc(ResearchSession.created_at)).offset(skip).limit(limit).all()
    return [
        {
            "id": s.id, "title": s.title, "query": s.query,
            "research_type": s.research_type, "jurisdiction": s.jurisdiction,
            "matter_id": s.matter_id,
            "created_at": s.created_at.isoformat(),
        }
        for s in sessions
    ]


@app.get("/api/research/precedents")
def search_precedents(
    q: Optional[str] = None, category: Optional[str] = None,
    skip: int = 0, limit: int = 20,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    query = db.query(LegalPrecedent).filter(LegalPrecedent.is_active == True)
    if q:
        query = query.filter(or_(
            LegalPrecedent.title.ilike(f"%{q}%"),
            LegalPrecedent.keywords.ilike(f"%{q}%"),
            LegalPrecedent.citation.ilike(f"%{q}%"),
        ))
    if category:
        query = query.filter(LegalPrecedent.category == category)
    total = query.count()
    results = query.offset(skip).limit(limit).all()
    return {
        "total": total,
        "items": [
            {
                "id": p.id, "title": p.title, "citation": p.citation,
                "court": p.court, "date": str(p.date) if p.date else None,
                "category": p.category, "source_type": p.source_type,
                "summary": p.summary,
            }
            for p in results
        ]
    }


# ─────────────────────────────────────────────────────────────
# SETTINGS
# ─────────────────────────────────────────────────────────────

@app.get("/api/settings")
def get_settings(db: Session = Depends(get_db),
                 current_user: User = Depends(get_current_active_user)):
    s = db.query(SystemSettings).first()
    if not s:
        return {}
    return {
        "firm_name": s.firm_name, "firm_address": s.firm_address,
        "firm_phone": s.firm_phone, "firm_email": s.firm_email,
        "default_currency": s.default_currency,
        "default_hourly_rate": s.default_hourly_rate,
        "vat_rate": s.vat_rate, "invoice_prefix": s.invoice_prefix,
        "trust_bank_name": s.trust_bank_name,
        "trust_account_number": s.trust_account_number,
        "operating_bank_name": s.operating_bank_name,
    }


@app.put("/api/settings")
def update_settings(req: SettingsUpdate, db: Session = Depends(get_db),
                    current_user: User = Depends(require_admin)):
    s = db.query(SystemSettings).first()
    if not s:
        s = SystemSettings()
        db.add(s)
    for k, v in req.dict(exclude_unset=True).items():
        setattr(s, k, v)
    s.updated_at = datetime.utcnow()
    db.commit()
    return {"message": "Settings updated"}


# ─────────────────────────────────────────────────────────────
# ADMIN
# ─────────────────────────────────────────────────────────────

@app.get("/api/admin/users")
def admin_users(db: Session = Depends(get_db),
                current_user: User = Depends(require_admin)):
    users = db.query(User).all()
    return [
        {
            "id": u.id, "username": u.username, "full_name": u.full_name,
            "email": u.email, "user_role": u.user_role,
            "is_active": u.is_active, "is_admin": u.is_admin,
            "title": u.title,
            "last_activity": u.last_activity.isoformat() if u.last_activity else None,
        }
        for u in users
    ]


@app.get("/api/admin/activity-log")
def activity_log(skip: int = 0, limit: int = 50,
                 db: Session = Depends(get_db),
                 current_user: User = Depends(require_admin)):
    activities = db.query(Activity).order_by(
        desc(Activity.timestamp)).offset(skip).limit(limit).all()
    return [
        {
            "id": a.id, "action": a.action,
            "description": a.description,
            "user_id": a.user_id,
            "resource_type": a.resource_type, "resource_id": a.resource_id,
            "timestamp": a.timestamp.isoformat(),
        }
        for a in activities
    ]


# ─────────────────────────────────────────────────────────────
# FRONTEND SERVING (PRODUCTION)
# ─────────────────────────────────────────────────────────────

def setup_frontend_serving():
    if getattr(sys, "frozen", False):
        frontend_dir = Path(sys._MEIPASS) / "frontend"
    else:
        frontend_dir = Path(__file__).parents[2] / "frontend" / "dist"

    if frontend_dir.exists():
        app.mount("/", StaticFiles(directory=str(frontend_dir), html=True), name="frontend")
        logger.info(f"Frontend served from: {frontend_dir}")
    else:
        logger.warning(f"Frontend dist not found at {frontend_dir}")
